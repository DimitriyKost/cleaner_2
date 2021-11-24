from docx import Document
import PySimpleGUI as Sg
import sys

Sg.theme('DarkGrey2')


def pars(file_docx):
    # переводим вордовский файл в формат HTML
    txt_r = "<p style='text-align: center;'><span style='color: #ff0000;'><strong>На мобильных устройствах расписание \
    можно двигать как по вертикали, так и по горизонтали.</strong></span></p>\
    <div style='max-width: 100%; overflow: scroll;'>\n\
    <table border='1'>\n\t<tbody>\n\t\t<tr>\n\
    \t\t\t<td style='text-align: center;' colspan='2' align='center' width='10%'><span style='font-size: 14pt;'>Дата \
    </span></td>\n\
    \t\t\t<td style='text-align: center;' width='50%'><span style='font-size: 14pt;'>Память святого или события</span> \
    </td>\n\
    \t\t\t<td style='text-align: center;' width='20%'><span style='font-size: 14pt;'>Утро</span></td>\n\
    \t\t\t<td style='text-align: center;' width='20%'><span style='font-size: 14pt;'>Вечер</span></td>\n\
    \t\t</tr>\n"
    txt_h = "\t\t\t<td style='text-align: center; font-size: 14pt;'>"
    # читаем файл и парсим его
    doc = Document(file_docx)
    table = doc.tables[0]
    for row in table.rows:  # просматриваем строки таблицы
        txt_r += "\t\t<tr>\n"
        for cell in row.cells:  # просматриваем ячейки строки
            txt_r += txt_h
            for par in cell.paragraphs:  # просматриваем абзацы ячейки
                txt_run = ""
                for run in par.runs:  # просматриваем части абзаца
                    color_d = str(run.font.color.rgb)
                    if color_d == "None" or color_d == "000000":
                        txt_run0 = run.text
                    else:
                        txt_run0 = "<span style='color:#" + color_d + ";'>" + run.text + "</span>"
                    txt_run += txt_run0
                txt_r += txt_run
                txt_r += "<br>"
            txt_r += "</td>\n"
        txt_r += "\t\t</tr>\n"
    txt_r += "\t</tbody>\n</table>\n</div>"
    return txt_r


layout = [[Sg.Text('Файл расписания: '), Sg.Text(size=(25, 1), key='-OUTPUT-')],
          [Sg.Button('Выбрать файл', size=(20, 2)), Sg.Button('Конвертировать', size=(20, 2))]]
window = Sg.Window('Чистильщик', layout)
while True:  # Event Loop
    event, values = window.read()
    if event == Sg.WIN_CLOSED or event == 'Exit':
        break
    if event == 'Выбрать файл':
        if len(sys.argv) == 1:
            f_name = Sg.popup_get_file('Выберите файл', file_types=(("Документ WORD", ".docx"),))
        else:
            f_name = sys.argv[1]
        window['-OUTPUT-'].update(f_name)
        if not f_name:
            Sg.popup("Cancel", "Не выбрали файл")
            raise SystemExit("Cancelling: no filename supplied")
        # fileR = f_name  # Берем название и путь файла для обработки
    if event == 'Конвертировать':
        text_t = pars(f_name)
        Sg.popup_scrolled(text_t, title="Текст после очистки", size=(120, 40))
        Sg.clipboard_set(text_t)
        Sg.popup("Очищенный текст скопирован в буфер обмена")
window.close()
